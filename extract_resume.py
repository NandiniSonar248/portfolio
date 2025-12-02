from docx import Document

# Open the Word document
doc = Document('Nandini Jitendra Sonar Final Resume.docx')

print("=" * 80)
print("COMPLETE RESUME CONTENT EXTRACTION")
print("=" * 80)

# Extract all paragraphs
print("\n### PARAGRAPHS ###")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(para.text)

# Extract all tables
print("\n\n### TABLES ###")
for table_idx, table in enumerate(doc.tables):
    print(f"\n--- Table {table_idx + 1} ---")
    for row_idx, row in enumerate(table.rows):
        cells_text = []
        for cell in row.cells:
            cells_text.append(cell.text.strip())
        print(" | ".join(cells_text))

print("\n" + "=" * 80)
print("END OF RESUME CONTENT")
print("=" * 80)
